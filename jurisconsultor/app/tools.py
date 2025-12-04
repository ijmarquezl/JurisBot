import requests
import json
import os
import re
import docx
from datetime import datetime
import logging
from jose import jwt


logger = logging.getLogger(__name__)

# This module defines the core logic for the tools the AI agent can use.

API_BASE_URL = os.getenv("BACKEND_API_URL", "http://jurisbot-project-manager-mcp:8000")
TEMPLATE_DIR = "../formatos/"
GENERATED_DOCS_PATH = "../documentos_generados/"

_auth_token = None
_tenant_id = None

def set_auth_token(token: str):
    """Sets the authentication token for the API calls for the current turn."""
    global _auth_token
    _auth_token = token

def set_tenant_id(tenant_id: str):
    """Sets the tenant ID for the API calls for the current turn."""
    global _tenant_id
    _tenant_id = tenant_id

def _get_headers() -> dict:
    """Helper function to get authentication headers."""
    if not _auth_token:
        # This clear error message is crucial for the agent to understand the problem.
        raise ValueError("Authentication token not set. I cannot use tools that require API calls. I must inform the user about a potential login issue.")
    if not _tenant_id:
        raise ValueError("Tenant ID not set. I cannot use tools that require API calls.")
    return {
        "Authorization": f"Bearer {_auth_token}",
        "X-Tenant-ID": _tenant_id,
        "Content-Type": "application/json",
    }

def get_template_placeholders(template_name: str) -> str:
    """Reads a .docx template and extracts its placeholders in a robust manner."""
    try:
        template_path = os.path.join(TEMPLATE_DIR, template_name)
        if not os.path.exists(template_path):
            return json.dumps({"error": f"Template '{template_name}' not found."})
        
        doc = docx.Document(template_path)
        placeholders = set()
        
        # Regex to find placeholders, including those split across runs
        placeholder_regex = re.compile(r'{{(.*?)}}')

        # Function to extract placeholders from a text block (paragraph or cell)
        def extract_from_block(block):
            # Reassemble text from runs to handle split placeholders
            full_run_text = "".join(run.text for run in block.runs)
            for match in placeholder_regex.finditer(full_run_text):
                placeholders.add(match.group(1).strip())

        # Process paragraphs
        for para in doc.paragraphs:
            extract_from_block(para)

        # Process tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    extract_from_block(cell) # Cells also have .runs
                    # Also check cell's paragraphs for good measure
                    for para_in_cell in cell.paragraphs:
                         extract_from_block(para_in_cell)

        if not placeholders:
            # If the above fails, fall back to a simpler full-text search
            full_text = "\n".join([p.text for p in doc.paragraphs])
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        full_text += "\n" + cell.text
            
            found = set(re.findall(r"{{(.*?)}}", full_text))
            if not found:
                return json.dumps({"error": f"No placeholders like '{{field}}' found in the template '{template_name}'."})
            return json.dumps([p.strip() for p in found])

        return json.dumps(list(placeholders))
    except Exception as e:
        return json.dumps({"error": f"Failed to read template. {e}"})

def fill_template_and_save_document(template_name: str, document_name: str, context: dict) -> str:
    """Fills and saves a .docx template with the provided context. Returns the path of the new file."""
    if not document_name or not context:
        return json.dumps({"error": "Called with missing document_name or context."})
    try:
        template_path = os.path.join(TEMPLATE_DIR, template_name)
        if not os.path.exists(template_path):
            return json.dumps({"error": f"Template '{template_name}' not found."})

        doc = docx.Document(template_path)

        # Helper to replace text in a block (paragraph or cell) while preserving some formatting
        def replace_text_in_block(block, key, value):
            search_text = f"{{{{{key}}}}}"
            if search_text in block.text:
                # Build up the new text
                new_text = block.text.replace(search_text, str(value))
                # Clear existing runs and add new run with replaced text
                for run in block.runs:
                    run.clear()
                block.add_run(new_text)

        # Replace placeholders in paragraphs
        for para in doc.paragraphs:
            for key, value in context.items():
                replace_text_in_block(para, key, value)

        # Replace placeholders in tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        for key, value in context.items():
                            replace_text_in_block(para, key, value)
        
        # Save the new document
        new_file_name = f"{document_name.replace(' ', '_')}_{datetime.utcnow().strftime('%Y%m%d%H%M%S')}.docx"
        new_file_path = os.path.join(GENERATED_DOCS_PATH, new_file_name)
        os.makedirs(GENERATED_DOCS_PATH, exist_ok=True) # Ensure directory exists
        doc.save(new_file_path)
        
        return json.dumps({"success": True, "file_path": new_file_path})
    except Exception as e:
        return json.dumps({"error": f"Failed to fill and save document. {e}"})



def list_projects() -> str:
    """Lists all projects the user is a member of."""
    try:
        headers = _get_headers()
        logger.debug(f"Calling API to list projects at {API_BASE_URL}/tools/list_projects with tenant_id: {_tenant_id}")
        response = requests.get(f"{API_BASE_URL}/tools/list_projects?tenant_id={_tenant_id}", headers=headers)
        response.raise_for_status()
        json_response = response.json()
        logger.debug(f"API response for list_projects: {json_response}")
        return json.dumps(json_response)
    except Exception as e:
        logger.error(f"Failed to list projects. {e}")
        return json.dumps({"error": f"Failed to list projects. {e}"})

def create_project(project_name: str, project_description: str = None) -> str:
    """Creates a new project via the backend API."""
    try:
        headers = _get_headers()
        
        # Decode the token to get the user's email
        # This assumes the same secret and algorithm are used as in the main app's security module
        secret_key = os.getenv("SECRET_KEY")
        algorithm = os.getenv("ALGORITHM")
        if not secret_key or not algorithm:
            raise ValueError("SECRET_KEY and ALGORITHM environment variables must be set.")
            
        try:
            payload_data = jwt.decode(_auth_token, secret_key, algorithms=[algorithm])
            user_email = payload_data.get("sub")
            if not user_email:
                return json.dumps({"error": "Could not extract user email from token."})
        except jwt.JWTError as e:
            logger.error(f"JWT decoding error: {e}")
            return json.dumps({"error": f"Invalid token. {e}"})

        payload = {
            "project_name": project_name,
            "tenant_id": _tenant_id,
            "user_email": user_email, # Add the user's email to the payload
            "project_description": project_description
        }
        response = requests.post(f"{API_BASE_URL}/tools/create_project", headers=headers, json=payload)
        response.raise_for_status()
        project_data = response.json()
        return json.dumps({"success": True, "project_id": project_data["project_id"], "project_name": project_data["project_name"]})
    except Exception as e:
        return json.dumps({"error": f"Failed to create project. {e}"})

def create_new_task(project_id: str, title: str, description: str = None) -> str:
    """Creates a new task for a given project via the backend API."""
    try:
        headers = _get_headers()
        payload = {
            "project_id": project_id,
            "title": title,
            "tenant_id": _tenant_id,
            "description": description
        }
        response = requests.post(f"{API_BASE_URL}/tools/create_task", headers=headers, json=payload)
        response.raise_for_status()
        task_data = response.json()
        return json.dumps({"success": True, "task_id": task_data["task_id"], "task_title": task_data["task_title"]})
    except Exception as e:
        return json.dumps({"error": f"Failed to create task. {e}"})

def list_tasks_for_project(project_id: str) -> str:
    """Lists all tasks for a given project via the backend API."""
    try:
        headers = _get_headers()
        response = requests.get(f"{API_BASE_URL}/tools/list_tasks_for_project?project_id={project_id}&tenant_id={_tenant_id}", headers=headers)
        response.raise_for_status()
        return json.dumps(response.json())
    except Exception as e:
        return json.dumps({"error": f"Failed to list tasks for project {project_id}. {e}"})